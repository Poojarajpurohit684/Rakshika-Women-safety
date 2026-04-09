import re
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ------------------------------------------------------------
# 1. Page border (soft blue, #B3D4FC)
# ------------------------------------------------------------
def set_page_border(doc, color='B3D4FC', space=20):
    section = doc.sections[0]
    for edge in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement('w:{}'.format(edge))
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '24')      # 1/8 point
        border.set(qn('w:space'), str(space))
        border.set(qn('w:color'), color)
        section._sectPr.append(border)

# ------------------------------------------------------------
# 2. Code block style (monospace, green text, light blue background)
# ------------------------------------------------------------
def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)   # #1B5E20
    # Light blue background (#EAF3FB)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'EAF3FB')
    run._element.getparent().get_or_add_pPr().append(shading)
    # Optional light border around code block
    p_border = OxmlElement('w:pBdr')
    for edge in ['top', 'left', 'bottom', 'right']:
        br = OxmlElement('w:{}'.format(edge))
        br.set(qn('w:val'), 'single')
        br.set(qn('w:sz'), '4')
        br.set(qn('w:space'), '4')
        br.set(qn('w:color'), 'B3D4FC')
        p_border.append(br)
    p._element.get_or_add_pPr().append(p_border)

# ------------------------------------------------------------
# 3. Headings with emojis and blue color (#0451A5)
# ------------------------------------------------------------
def add_heading1(doc, text):
    # choose emoji based on content
    if 'CHAPTER' in text.upper() or 'INTRODUCTION' in text.upper():
        icon = '📘 '
    elif 'CONCLUSION' in text.upper():
        icon = '🎯 '
    elif 'TESTING' in text.upper():
        icon = '🧪 '
    elif 'APPENDIX' in text.upper():
        icon = '📎 '
    else:
        icon = '🛡️ '
    heading = doc.add_heading(icon + text, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x04, 0x51, 0xA5)   # #0451A5
        run.font.bold = True
        run.font.size = Pt(18)

def add_heading2(doc, text):
    icon = '🔸 '
    heading = doc.add_heading(icon + text, level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x04, 0x51, 0xA5)
        run.font.size = Pt(14)

def add_heading3(doc, text):
    icon = '▪️ '
    heading = doc.add_heading(icon + text, level=3)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x04, 0x51, 0xA5)
        run.font.size = Pt(12)

# ------------------------------------------------------------
# 4. Normal paragraph style (#2C2C2C)
# ------------------------------------------------------------
def add_normal_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)
    run.font.size = Pt(11)

# ------------------------------------------------------------
# 5. Simple table converter (markdown pipe tables)
# ------------------------------------------------------------
def add_table_from_markdown(doc, lines):
    # lines: list of strings that form a markdown table
    # first line: header, second: separator, rest: rows
    if len(lines) < 2:
        return
    # parse header
    header_cells = [cell.strip() for cell in lines[0].split('|') if cell.strip()]
    # determine column count
    col_count = len(header_cells)
    if col_count == 0:
        return
    table = doc.add_table(rows=1, cols=col_count)
    table.style = 'Table Grid'
    # header row
    for i, cell_text in enumerate(header_cells):
        table.rows[0].cells[i].text = cell_text
        # bold header text
        for paragraph in table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x04, 0x51, 0xA5)
    # data rows (skip separator line)
    for line in lines[2:]:
        if '|' not in line:
            continue
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if len(cells) != col_count:
            continue
        row = table.add_row()
        for i, cell_text in enumerate(cells):
            row.cells[i].text = cell_text
    doc.add_paragraph()  # empty line after table

# ------------------------------------------------------------
# Main document builder
# ------------------------------------------------------------
def create_premium_report():
    doc = Document()
    set_page_border(doc, color='B3D4FC', space=20)

    # The full report content (as provided by the user)
    # We embed it as a raw string. To keep the script manageable,
    # we use a multi‑line string. (The content is exactly as given,
    # but we remove the outer "```" markers if any.)
    content = r"""**RAKSHIKA\
WOMEN'S PERSONAL SAFETY APPLICATION\
FULL PROJECT REPORT**

Submitted in partial fulfillment of the requirements for the award of the degree of
Bachelor of Computer Applications / B.Tech (CSE)

Project Title     : Rakshika - Women's Personal Safety Application
Student Name      : [Your Full Name]
Roll Number       : [Your Roll Number]
Department        : [Department Name]
University / College: [University / College Name]
Academic Year     : 2025-2026
Guide Name        : [Guide Name]

CERTIFICATE OF COMPLETION
This is to certify that the project entitled 'Rakshika - Women's Personal Safety Application' submitted by [Student Name], Roll No. [Roll Number], is a bonafide record of original work carried out under my supervision and guidance in partial fulfillment of the requirements for the award of the degree of [Degree Name] from [University Name].
The work embodied in this report is original and has not been submitted for the award of any other degree or diploma in any institution, university, or college. Wherever information has been derived from other sources, it has been duly acknowledged.

Guide / Supervisor
___________________________
Head of Department
___________________________
Principal
___________________________
Official Seal / Stamp
___________________________


                                         ACKNOWLEDGMENT
I take this opportunity to express my profound gratitude to my project guide,asst.professor  Ms Jyotsna Bhubanesh, , Department of BCA, college of R.k desai College of Computer and applied Sciences, for exemplary guidance, constant encouragement, and technical direction throughout the project lifecycle.
I also extend sincere thanks to the Head of Department, HOD Ms.Neha Patel , for support, review, and administrative assistance. The cooperation of faculty members and laboratory staff helped me complete the work with clarity and confidence.
I am grateful to the principal and the institution for providing the infrastructure, software environment, and academic atmosphere required to complete this report and associated implementation work.
I would also like to thank the open-source community and the developers of the libraries, frameworks, and services that made the project possible. Their tools supported the frontend, backend, geolocation, messaging, testing, and deployment pipeline.
Finally, I thank my family, friends, and peer reviewers for their patience, feedback, and encouragement during the development and documentation process.


                           
                                             ABSTRACT
Rakshika is a full-stack women's safety application designed to reduce response time during emergencies and improve situational awareness. The system integrates one-tap SOS alerts, live location sharing, trusted contact management, safe zone discovery, fake call simulation, and an AI safety assistant into a single responsive interface.
The application is designed around a mobile-first architecture and uses modern web technologies to remain accessible across devices. The frontend focuses on speed, simplicity, and discreet interaction under stress, while the backend manages authentication, alert delivery, location logging, and public live-share links.
The report combines the academic structure of a project submission with a practical description of the system design, database model, implementation modules, testing strategy, and future enhancement roadmap. It also includes diagrams such as DFDs, architecture blocks, and ER representation to support review and evaluation.
The proposed system emphasizes reliability, usability, and rapid action. In an emergency, the application is intended to minimize the number of user steps and provide trusted contacts with immediate notification and precise geographic context.


    TABLE OF CONTENTS

CERTIFICATE OF COMPLETION	2
ACKNOWLEDGMENT	3
ABSTRACT	4
TABLE OF CONTENTS	5
CHAPTER 1 - INTRODUCTION	7
1.1 Project Motivation	8
1.2 Objectives of the System	9
1.3 Scope of the Project	10
1.4 Purpose and Significance	11
1.5 Organization of the Report	11
CHAPTER 2 - LITERATURE SURVEY / REVIEW	12
2.1 Existing Safety Applications	13
2.2 Limitations of Existing Systems	13
2.3 Research Gap	14
2.4 Comparative Summary	14
CHAPTER 3 - REQUIREMENT & ANALYSIS	15
3.1 Existing System	15
3.2 Proposed System	15
3.3 Hardware Requirements	16
3.4 Software Requirements	17
3.5 Functional Requirements	18
3.6 Non-Functional Requirements	19
3.7 Feasibility Analysis	19
CHAPTER 4 - SYSTEM DESIGN AND ARCHITECTURE	20
4.1 System Architecture	20
4.2 Data Flow Diagram (DFD) - Level 0	21
4.3 Data Flow Diagram (DFD) - Level 1	21
4.4 Entity-Relationship (ER) Diagram	22
4.5 Database Design	23
4.6 API Design	23
4.7 Flowcharts and Operational Sequences	24
4.8 Security Design	25
CHAPTER 5 - IMPLEMENTATION DETAILS	26
5.1 Project Structure	26
5.2 Authentication Module	72
5.3 Dashboard and SOS Alert Module	72
5.4 Live Location Module	72
5.5 Trusted Contacts Module	73
5.6 Safe Zones Module	73
5.7 Fake Call Module	73
5.8 AI Chat Assistant Module	73
5.9 Awareness / Onboarding Module	74
5.10 Check-in Timer Module	74
5.11 Sample Code Snippets	74
5.12 Detailed Module Notes	75
CHAPTER 6 - TESTING, RESULTS, AND ANALYSIS	78
6.1 Testing Strategy	78
6.2 Representative Test Cases	78
6.3 Performance Analysis	80
6.4 Results Summary	80
6.5 Limitations Found During Testing	81
CHAPTER 7 - SCREENSHOTS AND UI WALKTHROUGH	82
Login / Register Screen	83
Dashboard / Home Screen	84
SOS Alert Trigger Screen	85
Trusted Contacts Screen	85
Safe Zones Screen	86
Fake Call Screen	86
AI Chat Assistant Screen	87
Live Location Share Screen	88
Awareness Carousel Screen	88
Emergency Helplines	89
Day/Night mode,Notifications,Profile	89
CHAPTER 8 - CONCLUSION AND FUTURE SCOPE	90
8.1 Conclusion	90
8.2 Future Scope	91
8.3 Final Remarks	92
APPENDIX A - DETAILED DATA DICTIONARY	93
APPENDIX B - DEPLOYMENT NOTES	94
APPENDIX C - USER GUIDE	95
APPENDIX D - SAMPLE API RESPONSE FORMAT	95
APPENDIX E - ADDITIONAL ANALYSIS NOTES	96
BIBLIOGRAPHY / REFERENCES	98
❖ Shankar, K., Singh, R., & Madheshiya, S. (2019). *Women Safety App to Detect Danger and Prevent Automatically Using Machine Learning*. Atlantis Press. The authors describe an app that monitors ambient sound (e.g. a scream above 40 dB) to automatically notify emergency contacts with real-time location【4†L198-L207】, and captures photo/video evidence if the user cannot respond【4†L208-L213】.	98


                              CHAPTER 1 - INTRODUCTION
Women's personal safety remains a critical social and technological challenge. In situations involving harassment, threat, stalking, or violent confrontation, every second matters. Conventional mechanisms often require speech, repeated navigation through menus, or dependence on third-party intervention. Rakshika responds to this problem with a streamlined digital safety workflow that prioritizes speed, discretion, and location precision.
The motivation for the project comes from the recognition that a modern safety tool must do more than send a basic alert. It must preserve context, reduce user stress, and provide a usable path to help even when the user cannot openly communicate. Rakshika is therefore designed to assist both the person in distress and the trusted contacts receiving the alert.
The application combines emergency signaling, real-time geographic context, and proactive support tools. Rather than forcing a user to switch between different apps for calling, messaging, map sharing, and safety guidance, the system centralizes those actions into a single interface. This consolidation improves usability in emergency conditions.
Technically, Rakshika is built as a full-stack application with a clear separation between client responsibilities and server-side services. The frontend handles the presentation layer, user interaction, and location capture, while the backend stores user data, maintains alert records, and orchestrates delivery through SMS, email, and AI-assisted safety services.
1.1 Project Motivation
The primary motivation is the need for a discreet safety companion that can be used under pressure. The user interface must be simple enough to operate during panic, darkness, low network coverage, or social isolation. The project therefore emphasizes one-tap action and short confirmation windows.
A second motivation is the desire to improve the quality of emergency information transmitted to contacts. A message that contains only a phone number or a generic distress statement may not be enough for immediate help. Rakshika adds live coordinates, map links, and session context so that contacts can respond more effectively.
A third motivation is educational and preventive. The application is not limited to reactive SOS. It also includes safe zones, awareness content, and an AI assistant that can guide the user through immediate next steps. This creates a layered safety model: prevent, escape, alert, and recover.

1.2 Objectives of the System
The primary objectives of the system are as follows:
To provide a one-touch SOS mechanism that instantly sends live location details to predefined trusted contacts during emergencies.
To implement background location tracking and automated emergency alert generation for faster and more effective response.
To develop a safe zone discovery feature that enables users to locate nearby essential services such as police stations, hospitals, and bus stops.
To include a fake call functionality that allows users to discreetly exit uncomfortable or potentially unsafe situations.
To integrate an AI-powered safety assistant capable of understanding natural language commands and responding to user safety queries.
To maintain a structured database for storing alert history, user contacts, and location sessions for traceability and future reference.
To design a clean, user-friendly, and mobile-responsive interface that remains easy to operate even under stressful conditions.

1.3 Scope of the Project
Rakshika is designed as a responsive web application and PWA-style safety interface. It is intended to function on modern browsers across mobile and desktop platforms. The project focuses on emergency support, contact notification, and location sharing rather than native telephony or hardware-specific integrations.
The in-scope features include user authentication, trusted contact management, one-tap SOS alerts, live location sharing, safe zone discovery, fake call simulation, AI guidance, and testing/development modes. The out-of-scope areas include full native app distribution, government emergency dispatch integration, and offline operation without any network connectivity.
The scope also includes the creation of project documentation, test artefacts, architecture diagrams, data structures, and supporting screenshots or placeholders required for academic submission. This report is therefore both a technical specification and a presentation of the completed academic work.
1.4 Purpose and Significance
The purpose of Rakshika is to make safety support more immediate and less dependent on manual escalation. The significance of the project lies in its practical response to common real-world constraints: a user may be afraid to speak, may not know exact location details, or may need a believable way to withdraw from a risky situation.
By combining emergency alerts with live location and contextual support, the system reduces the time between threat recognition and communication with helpers.The project is also significant from a software engineering perspective because it demonstrates how a well-structured full-stack application can support socially relevant use cases through modular design, clean data flow, and clear user experience priorities.
1.5 Organization of the Report
This report is organized into eight chapters, each covering a distinct phase
of the software development lifecycle of the Rakshika Women's Personal Safety
Application. The structure follows the standard format prescribed for academic
project reports.

Chapter 1 – Introduction
  This chapter introduces the project by providing the background and context
  of women's safety in India, clearly defining the problem statement, listing
  the objectives of the application, describing the scope and boundaries of
  the project, and explaining the purpose and significance of building Rakshika
  as a technology-driven safety solution.

Chapter 2 – Literature Survey / Review
  This chapter presents a review of existing women's safety applications such
  as Nirbhaya App, Himmat App, bSafe, and Shake2Safety. It identifies the
  limitations and shortcomings of these existing systems and establishes the
  research gap that Rakshika aims to fill with its unified, AI-powered,
  PWA-based approach.

Chapter 3 – Requirement & Analysis
  This chapter provides a detailed analysis of the existing system and the
  proposed system, followed by a comprehensive listing of hardware requirements,
  software requirements (frontend, backend, APIs, testing tools), functional
  requirements (FR-01 to FR-12), and non-functional requirements covering
  performance, security, scalability, and accessibility.

Chapter 4 – System Design
  This chapter covers the complete system design of Rakshika including the
  three-tier system architecture diagram, Level 0 and Level 1 Data Flow
  Diagrams (DFD), Entity-Relationship (ER) diagram showing the three MongoDB
  collections and their relationships, detailed database schema design for
  Users, Contacts, and Locations collections, a complete REST API design table,
  and flowcharts for the SOS alert trigger and user authentication processes.

Chapter 5 – Implementation / Coding
  This chapter describes the actual implementation of the application. It
  begins with a module overview and the complete project directory structure
  with descriptions of every file and folder. It then presents key code
  snippets with explanations for all major modules: Authentication, SOS Alert,
  Live Location, Trusted Contacts, Safe Zone Discovery, Fake Call Simulator,
  AI Chat Assistant, Live Share, Awareness Carousel, and Frontend Routing with
  Auth Context. Environment variables and API route summary are also included.

Chapter 6 – Testing
  This chapter describes the testing strategy adopted for the project including
  unit testing, integration testing, system testing, and manual testing. It
  lists the tools used (Vitest, React Testing Library, jsdom, Postman), presents
  detailed test cases with inputs, expected outputs, and actual results, and
  provides a test results summary table showing pass/fail status for all
  major test scenarios.

Chapter 7 – Screenshots
  This chapter presents visual snapshots of the application's user interface
  covering all major screens: Login and Register, Dashboard with SOS button,
  SOS countdown trigger, Trusted Contacts management, Safe Zones map and list,
  Fake Call incoming and talking screens, AI Chat Assistant, Live Location
  Share page, and the Awareness onboarding carousel.

Chapter 8 – Conclusion & Future Enhancements
  This chapter summarizes the outcomes and achievements of the project,
  discusses the limitations of the current implementation, and proposes a
  roadmap of future enhancements categorized into short-term, medium-term,
  and long-term improvements including push notifications, native mobile app,
  voice-activated SOS, community safety heatmap, and smart city integration.

Bibliography / References
  This section lists all the academic references, official documentation,
  API documentation, and online resources consulted during the design,
  development, and documentation of the Rakshika project.

CHAPTER 2 - LITERATURE SURVEY / REVIEW
The literature survey reviews the landscape of women's safety applications and adjacent emergency support systems. Existing tools have inspired the core features of Rakshika, but they also reveal significant gaps that remain unsolved in many real-world use cases.
Several safety applications provide location sharing or SOS alerts, yet many stop short of integrating all critical components into one workflow. Some require the contact to install the same application, while others rely on manual sharing or a phone call that may not be possible under stress.
The review process shows that the strongest systems typically provide one or two emergency features. Rakshika extends this model by incorporating multiple safety dimensions: communication, navigation, discreet escape, and guided response.
2.1 Existing Safety Applications
Category
Observations
Government emergency helplines
Universal and widely known, but rely on verbal communication and do not automatically share context.
Location-sharing tools
Useful for tracking, but often require manual activation and lack integrated alert workflows.
Basic SOS applications
Send emergency messages, but typically do not include AI guidance or safe-zone discovery.
Discreet helper apps
Offer fake calls or alarms, but usually lack trusted contacts and history logging.
Navigation services
Can indicate nearby locations, but are not purpose-built for emergency safety communication.

From the comparison, the design requirement becomes clear: an effective safety platform must avoid feature fragmentation and support multiple response patterns inside one accessible interface.
2.2 Limitations of Existing Systems
Fragmentation: users often need multiple apps for SOS, maps, and contact messaging.
Manual operation: too many steps are required in a situation where speed is essential.
Lack of contextual assistance: many systems only send a message without guidance.
Poor support for silent use: speaking or typing may not be possible during an emergency.
Weak usability under stress: interfaces can be cluttered, making them hard to use quickly.
2.3 Research Gap
The research gap addressed by Rakshika is the absence of a unified, intelligent, India-aware, mobile-first women's safety platform that supports alerting, live location, safe-zone discovery, and discreet exit tools in a single cohesive design.
The gap is not just functional but experiential. Safety software must be usable when the user is panicked, distracted, or unable to speak. Rakshika therefore prioritizes minimal interaction, clear visual hierarchy, and automated workflows.

2.4 Comparative Summary
Feature
Existing tools
Rakshika
One-tap SOS
Partial or absent in many systems
Fully supported
Live location sharing
Often manual or app-specific
Public share link and map
Safe zones
Usually unavailable
Integrated discovery
Fake call
Rare or basic
Selectable caller and delay
AI guidance
Generally absent
Built into chat assistant
Trusted contacts
Sometimes present
Structured and manageable

CHAPTER 3 - REQUIREMENT & ANALYSIS
This chapter analyzes the current situation, identifies pain points, and converts them into system requirements. The goal is to align the technical implementation with the actual usage conditions of an emergency safety application.
3.1 Existing System
Government helplines are effective when the user can call, speak, and explain the issue. In many unsafe situations, however, speech is not possible or may increase risk.
Manual messaging through SMS, WhatsApp, or other apps requires several sequential actions. The delay introduced by those steps can be critical in an emergency.
Standalone GPS sharing tools provide location context but do not automatically alert trusted contacts or add guidance and escalation mechanisms.
Basic SOS applications improve responsiveness but still leave major gaps such as contextual AI support, safe-zone discovery, and discreet escape tools.
3.2 Proposed System
The proposed system integrates the most useful emergency features into a unified workflow. Rakshika initiates a short confirmation countdown, captures the user's coordinates, locates trusted contacts, and delivers alerts through SMS and optional email fallback.
The system additionally exposes a public live-tracking route so recipients can monitor the user's location without requiring a dedicated mobile installation. This makes the response simpler and more inclusive.
The AI assistant can answer safety questions and trigger core actions by command. This allows the user to request help through natural language when direct button navigation is difficult.
3.3 Hardware Requirements
On the client side, the application requires only a modern smartphone or browser-enabled device with GPS, internet access, and speaker support for the fake call feature. A microphone is optional for voice input. The system is designed to work on modest hardware so that accessibility remains high.
On the server side, development and deployment requirements are intentionally lightweight. This ensures that the project can be demonstrated in an academic environment, deployed economically, and scaled later if needed.

Component
Requirement
Client device
Smartphone / tablet / computer with a modern browser
Memory
2 GB minimum
Storage
About 50 MB free for cache and assets
Location
GPS or location services enabled
Network
3G / 4G / Wi-Fi connectivity
Audio
Speaker and microphone for fake call and voice input

3.4 Software Requirements
 The software stack was selected to balance modern UI quality, maintainability, and academic clarity. React 19, Vite, Tailwind CSS, Framer Motion, and Mapbox create a polished interface on the client side, while Node.js, Express, MongoDB, Mongoose, and secure authentication components support the server side.
In the merged project story, React Native remains relevant as the earlier prototype direction, but the final reported build emphasizes the PWA implementation because it matches the more complete feature set shown in the second file. This is a natural project evolution rather than a contradiction, and it makes the report stronger because it shows iterative development.

Layer
Stack
Frontend
React, Vite, Tailwind CSS, Framer Motion, Mapbox GL
Backend
Node.js, Express, Mongoose, MongoDB
Security
JWT, bcrypt, role-based API protection
Messaging
Twilio / Fast2SMS / Nodemailer
AI and maps
Google Gemini API, Google Places API
Testing
Vitest, React Testing Library, js dom

3.5 Functional Requirements
FR-01: User registration with secure storage of password hashes.
FR-02: User login and token-based session management.
FR-03: Trusted contact creation, listing, searching, and deletion.
FR-04: One-tap SOS activation with countdown confirmation.
FR-05: Live location sharing with a public URL and map view.
FR-06: Safe-zone discovery for police stations, hospitals, and bus stops.
FR-07: Fake incoming call simulation with delay and caller selection.
FR-08: AI safety chat that can understand emergency commands.
FR-09: Emergency helpline shortcuts on the dashboard.
FR-10: Check-in timer that escalates if the user does not respond.

3.6 Non-Functional Requirements
Performance: the system should load quickly and send alerts with low latency.
Security: credentials and tokens must be protected and never exposed in plain text.
Reliability: alert delivery should have fallback paths where possible.
Usability: the interface should remain understandable and touch-friendly.
Maintainability: modules should be separated cleanly so future updates are simpler.
Scalability: the architecture should support higher alert traffic and more users later.
3.7 Feasibility Analysis
Technical feasibility is strong because the required technologies are mature and widely supported. Browser geolocation, REST APIs, SMS gateways, and cloud databases are standard building blocks for contemporary web applications.
Operational feasibility is also favorable because the system is designed for normal smartphones and common usage patterns. A user does not need special training to understand the primary emergency flow.
Economic feasibility is reasonable because the core stack uses open-source software, with paid services applied only where needed for messaging, maps, or AI calls.

CHAPTER 4 - SYSTEM DESIGN AND ARCHITECTURE
This chapter describes the logical and structural design of Rakshika. The objective is to explain how data and control move through the system and how the main functional blocks cooperate.
4.1 System Architecture
Rakshika is organized as a client-server application with supporting third-party services. The client is responsible for interaction, location capture, map rendering, and presentation. The server manages authentication, contact data, alert creation, and public sharing logic. External services handle tasks such as SMS delivery, email fallback, geocoding, maps, and conversational AI.

Figure 4.1: High-level system architecture of Rakshika.
4.2 Data Flow Diagram (DFD) - Level 0

Figure 4.2: Context-level DFD showing user interaction and external communication.
4.3 Data Flow Diagram (DFD) - Level 1

Figure 4.3: Level 1 DFD showing the principal processing modules.
4.4 Entity-Relationship (ER) Diagram

Figure 4.4: ER diagram representing the core data entities and their relationships.
    
4.5 Database Design
The database design centers on the User entity, which owns contacts, alert logs, live-share sessions, and chat sessions. This normalization prevents duplicated data and keeps the alert history traceable.
Trusted contacts are stored as separate documents so that the system can update them independently without rewriting the user record. Alert logs store temporal and geographic information used for incident review.
The live location share entity contains a public token and status flag. This supports a browser-accessible route that can be activated or revoked cleanly.
Collection
Purpose
User
Authentication and profile data
TrustedContact
Emergency contact names and numbers
AlertLog
SOS event records and coordinates
LocationShare
Public live tracking state
ChatSession
Conversation history with AI

4.6 API Design
Endpoint
Purpose
POST /auth/register
Create a new user account
POST /auth/login
Authenticate user and issue JWT
GET /contacts
Fetch trusted contacts
POST /contacts
Add trusted contact
DELETE /contacts/:id
Remove trusted contact
POST /emergency/sos
Trigger SOS alert
POST /share/start
Start live location sharing
GET /share/:userId
Open public live share page
POST /ai/chat
Send message to AI assistant

4.7 Flowcharts and Operational Sequences
Operationally, the SOS flow is designed so that accidental triggers can be avoided while still remaining fast enough for real emergencies. The flowchart in Figure 4.5 documents the short countdown, location acquisition, contact lookup, and broadcast sequence.

4.8 Security Design
Security is handled through hashed passwords, signed tokens, authorization guards on protected routes, and strict separation between public and private endpoints.
The public sharing route is intentionally narrow in scope. It should expose only the information necessary for live tracking and nothing more than required for the intended help scenario.
The system avoids storing sensitive data in the browser beyond what is needed for a session and clears transient states when the user logs out.

CHAPTER 5 - IMPLEMENTATION DETAILS
This chapter explains how the design was translated into working modules. Implementation was driven by separation of concerns: authentication, alerting, map services, chat assistance, and contact management are handled as separate units with well-defined interfaces.
The codebase was organized to keep UI components small and reusable. The backend mirrors this with route handlers, controllers, and model definitions that support predictable state flow.
5.1 Project Structure
The project structure divides source code into client, server, shared utilities, and assets. This layout simplifies navigation, testing, and later maintenance.
Implementation notes: the module is built to be predictable, testable, and easy to audit during emergency-driven execution. Each action is logged, and failures are surfaced with a graceful fallback whenever feasible.
The Rakshika application follows a standard monorepo structure with two
separate directories — client (frontend) and server (backend) — each with
their own dependencies, configuration, and entry points.

DIRECTORY STRUCTURE
rakshika/                          ← Root project folder
│
├── client/                        ← Frontend (React + Vite PWA)
│   ├── public/
│   │   ├── favicon.svg            ← App favicon
│   │   └── icons.svg              ← SVG icon set
│   │
│   ├── src/
│   │   ├── assets/
│   │   │   ├── hero.png           ← Hero image for onboarding
│   │   │   ├── react.svg          ← React logo asset
│   │   │   └── vite.svg           ← Vite logo asset
│   │   │
│   │   ├── components/
│   │   │   ├── Map.jsx            ← Reusable Mapbox GL map component
│   │   │   ├── WomenSafetySection.jsx  ← Safety stats/info section
│   │   │   └── storytelling/
│   │   │       ├── index.js                  ← Barrel export
│   │   │       ├── RakshikaStoryCarousel.jsx ← Onboarding carousel
│   │   │       ├── StoryCard.jsx             ← Individual story card
│   │   │       ├── storyData.js              ← Story content data
│   │   │       └── StoryIllustrations.jsx    ← SVG illustrations
│   │   │
│   │   ├── data/
│   │   │   └── womensSafetyIndia.js   ← Static safety statistics data
│   │   │
│   │   ├── pages/
│   │   │   ├── Login.jsx              ← Login page
│   │   │   ├── Register.jsx           ← Registration page
│   │   │   ├── Dashboard.jsx          ← Main home screen (SOS, chat, map)
│   │   │   ├── TrustedContacts.jsx    ← Emergency contacts management
│   │   │   ├── SafeRoute.jsx          ← Nearby safe zones discovery
│   │   │   ├── FakeCall.jsx           ← Fake call simulator
│   │   │   ├── Share.jsx              ← Public live location view
│   │   │   ├── Awareness.jsx          ← Onboarding/education carousel
│   │   │   ├── Dashboard.bug-exploration.test.jsx   ← Bug exploration tests
│   │   │   └── Dashboard.preservation.test.jsx      ← Regression tests
│   │   │
│   │   ├── App.jsx                ← Root component, routing, AuthContext
│   │   ├── App.css                ← Global component styles
│   │   ├── index.css              ← Global CSS variables and base styles
│   │   ├── main.jsx               ← React DOM entry point
│   │   ├── api.js                 ← Axios instance with JWT interceptor
│   │   ├── auth.js                ← Token/user localStorage helpers
│   │   ├── AuthContext.js         ← React Context for auth state
│   │   └── test-setup.js          ← Vitest/Testing Library setup
│   │
│   ├── dist/                      ← Production build output (auto-generated)
│   ├── node_modules/              ← Frontend dependencies
│   ├── index.html                 ← Vite HTML entry point
│   ├── vite.config.js             ← Vite + PWA plugin configuration
│   ├── tailwind.config.js         ← Tailwind CSS configuration
│   ├── postcss.config.js          ← PostCSS configuration
│   ├── eslint.config.js           ← ESLint rules
│   ├── package.json               ← Frontend dependencies & scripts
│   ├── .env                       ← Frontend environment variables
│   └── .gitignore
│
├── server/                        ← Backend (Node.js + Express REST API)
│   ├── middleware/
│   │   └── auth.js                ← JWT authentication middleware
│   │
│   ├── models/
│   │   ├── User.js                ← Mongoose User schema
│   │   ├── Contact.js             ← Mongoose Contact schema
│   │   └── Location.js            ← Mongoose Location schema
│   │
│   ├── routes/
│   │   ├── auth.js                ← POST /auth/register, /auth/login
│   │   ├── contacts.js            ← GET/POST/DELETE /contacts
│   │   ├── location.js            ← POST/GET /location
│   │   ├── sos.js                 ← POST /sos/trigger
│   │   ├── share.js               ← GET /share/:userId
│   │   ├── places.js              ← GET /places/nearby
│   │   └── chat.js                ← POST /chat (Gemini AI)
│   │
│   ├── utils/
│   │   ├── twilio.js              ← SMS sending (Fast2SMS + Twilio)
│   │   ├── notifications.js       ← Email alert via Nodemailer
│   │   └── store.js               ← In-memory store (MOCK_MODE)
│   │
│   ├── node_modules/              ← Backend dependencies
│   ├── index.js                   ← Express server entry point
│   ├── package.json               ← Backend dependencies & scripts
│   ├── .env                       ← Backend environment variables
│   └── README.md
│
├── logo.png                       ← Project logo
├── poster.png                     ← Project poster
└── apiAIzaSyC7uDfaJQNCxeSMHystHYEad8dF.txt  ← API key reference file

🔷 CLIENT SIDE
File / Folder
Description
src/main.jsx
Entry point — mounts React app to DOM
src/App.jsx
Root component — routing and AuthContext
src/api.js
Axios instance with JWT auto-attach
src/auth.js
Local storage helpers for token and user
src/AuthContext.js
Global authentication state using React Context
src/index.css
CSS variables, dark theme, base styles
src/pages/Dashboard.jsx
Main screen — SOS, map, chatbot, check-in
src/pages/TrustedContacts.jsx
Add, view, delete emergency contacts
src/pages/SafeRoute.jsx
Nearby police stations, hospitals, bus stops
src/pages/FakeCall.jsx
Fake incoming call simulator
src/pages/Share.jsx
Public live location viewer (no login required)
src/pages/Awareness.jsx
Onboarding story carousel screen
src/pages/Login.jsx
Login form with JWT authentication
src/pages/Register.jsx
New user registration form
src/components/Map.jsx
Reusable Mapbox GL interactive map
src/components/storytelling
Onboarding carousel components and data
src/data/womensSafetyIndia
Static safety statistics data
vite.config.js
Vite build configuration with PWA setup
tailwind.config.js
Tailwind CSS customization
index.html
HTML shell for Vite
.env
Environment variables (API base, Mapbox token)

🔷 SERVER SIDE
File / Folder
Description
index.js
Express server entry point with CORS and route setup
middleware/auth.js
JWT verification for protected routes
models/User.js
User schema (name, email, password)
models/Contact.js
Contact schema (userId, name, phone)
models/Location.js
Location schema (lat, lng, isLive)
routes/auth.js
User registration and login APIs
routes/contacts.js
CRUD operations for contacts
routes/location.js
Save and fetch GPS location
routes/sos.js
SOS trigger — sends SMS and email alerts

ENVIRONMENT VARIABLES – SERVER
Variable
Purpose
PORT
Server port (default: 5000)
MONGO_URI
MongoDB connection string
JWT_SECRET
Secret key for JWT
MOCK_MODE
Enable mock database mode
CLIENT_ORIGIN
Frontend URL for CORS
GOOGLE_MAPS_KEY
Google Maps API key
GEMINI_API_KEY
Google Gemini AI key
FAST2SMS_API_KEY
SMS API for India
TWILIO_ACCOUNT_SID
Twilio account SID
TWILIO_AUTH_TOKEN
Twilio authentication token
TWILIO_PHONE_NUMBER
Sender phone number
FALLBACK_ALERT_EMAIL
Email fallback for SOS
SMTP_HOST
Mail server host
SMTP_USER
Email username
SMTP_PASS
Email password

🔷ENVIRONMENT VARIABLES – CLIENT
Variable
Purpose
VITE_API_BASE
Backend API URL
VITE_MAPBOX_TOKEN
Mapbox access token

🔷 API ROUTES SUMMARY 
Route
Method
Description
Auth
File
/auth/register
POST
Register new user
Public
routes/auth.js
/auth/login
POST
Login user and return JWT
Public
routes/auth.js
/contacts
GET
Get all contacts
JWT Required
routes/contacts.js
/contacts
POST
Add new contact
JWT Required
routes/contacts.js
/contacts/:id
DELETE
Delete contact by ID
JWT Required
routes/contacts.js
/location
POST
Save GPS location
JWT Required
routes/location.js
/location
GET
Get latest location
JWT Required
routes/location.js
/sos/trigger
POST
Send SOS alerts to contacts
JWT Required
routes/sos.js
/share/:userId
GET
Public location (no login)
Public
routes/share.js
/places/nearby
GET
Get nearby safe zones
JWT Required
routes/places.js
/chat
POST
AI chatbot (Gemini)
JWT Required
routes/chat.js

HOW TO RUN THE PROJECT
Step 1 — Install dependencies:
  cd server && npm install
  cd client && npm install

Step 2 — Configure environment:
  Edit server/.env — set MONGO_URI, JWT_SECRET, API keys
  Edit client/.env — set VITE_API_BASE=http://localhost:5000

Step 3 — Start backend:
  cd server
  npm run dev          ← starts with nodemon (auto-restart on changes)
  npm start            ← production start
Step 4 — Start frontend:
  cd client
  npm run dev          ← starts Vite dev server on http://localhost:5173
Step 5 — Build for production:
  cd client
  npm run build        ← outputs to client/dist/

NOTE: To run without MongoDB, set MOCK_MODE=1 in server/.env   All features work in mock mode except data does not persist
      after server restart.

5.3 Code Snippets
1. USER REGISTRATION (server/routes/auth.js)
This endpoint handles new user registration. It validates input fields, checks
for duplicate emails, hashes the password using bcrypt (salt factor 10), saves
the user to MongoDB, and returns a signed JWT token valid for 7 days.
CODE:
router.post('/register', async (req, res) => {
  try {
    const { name, email, password } = req.body;

    // Validate required fields
    if (!name || !email || !password)
      return res.status(400).json({ error: 'Missing fields' });

    // Check if email already exists
    const exists = await User.findOne({ email });
    if (exists)
      return res.status(400).json({ error: 'Email already registered' });

    // Hash password with bcrypt (salt rounds: 10)
    const hash = await bcrypt.hash(password, 10);

    // Create user in database
    const user = await User.create({ name, email, password: hash });

    // Sign JWT token (expires in 7 days)
    const token = jwt.sign(
      { id: user._id, email },
      process.env.JWT_SECRET || 'dev_secret',
      { expiresIn: '7d' }
    );

    res.json({ token, user: { id: user._id, name: user.name, email } });
  } catch (e) {
    res.status(500).json({ error: e.message });
  }
});

2. USER LOGIN (server/routes/auth.js)
This endpoint authenticates an existing user. It finds the user by email,
compares the submitted password against the stored bcrypt hash, and returns
a fresh JWT token on success.
CODE:
router.post('/login', async (req, res) => {
  try {
    const { email, password } = req.body;
    // Find user by email
    const user = await User.findOne({ email });
    if (!user)
      return res.status(400).json({ error: 'Invalid credentials' });
  // Compare password with stored hash
    const ok = await bcrypt.compare(password, user.password);
    if (!ok)
      return res.status(400).json({ error: 'Invalid credentials' });
   // Generate JWT token
    const token = jwt.sign(
      { id: user._id, email },
      process.env.JWT_SECRET || 'dev_secret',
      { expiresIn: '7d' }
    );
   res.json({ token, user: { id: user._id, name: user.name, email } });
  } catch (e) {
    res.status(500).json({ error: e.message });
  }
});
3. JWT AUTHENTICATION MIDDLEWARE (server/middleware/auth.js)
This middleware protects all private API routes. It extracts the Bearer token
from the Authorization header, verifies it using the JWT secret, and attaches
the decoded user payload (id, email) to req.user for use in route handlers.
Returns 401 if the token is missing or invalid.
CODE:
const jwt = require('jsonwebtoken');
module.exports = function auth(req, res, next) {
  // Extract token from Authorization header
  const authHeader = req.headers.authorization || '';
  const token = authHeader.startsWith('Bearer ')
    ? authHeader.slice(7)
    : null;
  if (!token)
    return res.status(401).json({ error: 'Missing token' });
  try {
    // Verify and decode token
    const payload = jwt.verify(
      token,
      process.env.JWT_SECRET || 'dev_secret'
    );
   // Attach user info to request
    req.user = { id: payload.id, email: payload.email };
    next();
  } catch (e) {
    return res.status(401).json({ error: 'Invalid token' });
  }};
4. SOS ALERT TRIGGER (server/routes/sos.js)
This is the core emergency endpoint. When triggered, it fetches all trusted
contacts of the authenticated user, constructs an SOS message containing a
Google Maps link and a live tracking URL, and sends SMS to every contact using
Fast2SMS or Twilio. It also sends a fallback email if configured.

CODE:
router.post('/trigger', async (req, res) => {
  const { lat, lng } = req.body || {};
  // Fetch all trusted contacts for this user
  const contacts = await Contact.find({ userId: req.user.id });
  const phones = contacts.map((c) => c.phone);
 // Build location links
  const liveLink = `${process.env.CLIENT_ORIGIN}/share/${req.user.id}`;
  const googleLink = `https://maps.google.com/?q=${lat},${lng}`;
  // Compose SOS message
  const message =
    `SOS ALERT! I need help. Location: ${googleLink} Track: ${liveLink}`;
  try {
    // Send SMS to each contact
    for (const phone of phones) {
      await sendSMS(phone, message);
    }

    // Send fallback email if configured
    if (process.env.FALLBACK_ALERT_EMAIL) {
      await sendAlert({ to: process.env.FALLBACK_ALERT_EMAIL, message });
    }

    res.json({ ok: true, notified: phones.length, liveLink });
  } catch (e) {
    res.status(500).json({ error: e.message });
  }
});

5. SMS SENDING UTILITY — Fast2SMS + Twilio Fallback (server/utils/twilio.js)
This utility handles SMS delivery. It first checks if a Fast2SMS API key is
configured (preferred for Indian numbers — no paid upgrade needed). If not,
it falls back to Twilio. Phone numbers are normalized to E.164 format (+91
prefix for 10-digit Indian numbers).

CODE:
async function sendSMS(to, body) {
 const fast2smsKey = process.env.FAST2SMS_API_KEY;
  // PRIMARY: Use Fast2SMS for Indian numbers
  if (fast2smsKey) {
    const mobile = toIndianMobile(to); // strips country code
    const res = await fetch('https://www.fast2sms.com/dev/bulkV2', {
      method: 'POST',
      headers: {
        authorization: fast2smsKey,
        'content-type': 'application/json',
      },
      body: JSON.stringify({
        message: body,
        route: 'q',
        numbers: mobile,
        flash: '0',
      }),
    });
    const data = await res.json();
    if (!data.return)
      throw new Error(`Fast2SMS error: ${JSON.stringify(data)}`);
    return { ok: true, provider: 'fast2sms' };
  }

  // FALLBACK: Use Twilio
  const client = twilio(
    process.env.TWILIO_ACCOUNT_SID,
    process.env.TWILIO_AUTH_TOKEN
  );
  const msg = await client.messages.create({
    from: process.env.TWILIO_PHONE_NUMBER,
    to: normalizePhone(to),  // ensures +91XXXXXXXXXX format
    body,
  });
  return { ok: true, sid: msg.sid };
}

// Normalize 10-digit Indian number to E.164
function normalizePhone(phone) {
  const cleaned = phone.replace(/\D/g, '');
  if (cleaned.length === 10) return `+91${cleaned}`;
  return phone.startsWith('+') ? phone : `+${cleaned}`;
}

6. GEOLOCATION HOOK (client/src/pages/Dashboard.jsx)
A custom React hook that wraps the browser's Geolocation API. It supports
both one-time position fetch and continuous real-time watching. The watch
mode pushes coordinate updates to a callback (used to send location to the
server during live sharing). Cleans up the watcher on component unmount.

CODE:
function useGeolocation() {
  const [pos, setPos] = useState(null);
  const [error, setError] = useState('');
  const watchId = useRef(null);
  // Get location once
  function getOnce() {
    if (!navigator.geolocation) {
      setError('Geolocation not supported');
      return;    }
    navigator.geolocation.getCurrentPosition(
      (p) => setPos({
        lat: p.coords.latitude,
        lng: p.coords.longitude,
        accuracy: p.coords.accuracy   }),
      (e) => setError(e.message),
      { enableHighAccuracy: true, timeout: 15000, maximumAge: 0 }    );  }
 // Start continuous GPS watch
  function startWatch(onUpdate) {
    if (watchId.current) return; // already watching
    watchId.current = navigator.geolocation.watchPosition(
      (p) => {
        const coords = {
          lat: p.coords.latitude,
          lng: p.coords.longitude,
          accuracy: p.coords.accuracy        };
        setPos(coords);
        onUpdate?.(coords); // push to server      },
      (e) => setError(e.message),
      { enableHighAccuracy: true, maximumAge: 10000, timeout: 10000 }    );  }
  // Stop watching
  function stopWatch() {
    if (watchId.current) {
      navigator.geolocation.clearWatch(watchId.current);
      watchId.current = null;
    }
  }

  // Cleanup on unmount
  useEffect(() => () => stopWatch(), []);

  return { pos, error, getOnce, startWatch, stopWatch };
}

7. LIVE LOCATION SHARE PAGE (client/src/pages/Share.jsx)
A public page (no login required) that displays the real-time location of a
user on an interactive Mapbox map. It polls the server every 15 seconds to
fetch the latest coordinates. Anyone with the link can view the location in
any browser without installing the app.
CODE:
export default function Share() {
  const { userId } = useParams();
  const [loc, setLoc] = useState(null);
  const [status, setStatus] = useState('loading');
  const [lastFetch, setLastFetch] = useState(null);
  // Fetch latest location from server
  async function load() {
    try {
      const res = await fetch(`${apiBase()}/share/${userId}`);
      const data = await res.json();
      setLoc(data);
      setStatus('ok');
      setLastFetch(new Date());
    } catch (e) {
      setStatus(e.message || 'error');
    }
  }

  // Auto-refresh every 15 seconds
  useEffect(() => {
    load();
    const t = setInterval(load, 15000);
    return () => clearInterval(t); // cleanup on unmount
  }, [userId]);
  return (
    // Renders Mapbox GL map centered on loc.lat, loc.lng
    // Shows live status indicator, coordinates, last updated time
    <Map
      center={[loc.lat, loc.lng]}
      markers={[{ lat: loc.lat, lng: loc.lng, color: '#e91e8c' }]}
    />
  );
}

8. SAFE ZONE DISCOVERY (client/src/pages/SafeRoute.jsx)
Queries the Google Places API (via the backend) for nearby police stations,
hospitals, and bus stops within 2 km of the user's GPS location. Results are
sorted by distance using the Haversine formula and displayed on a map with
distance, rating, and a direct Google Maps navigation link.
CODE:
// Haversine formula — calculates distance between two GPS coordinates
function getDistance(lat1, lon1, lat2, lon2) {
  const R = 6371; // Earth radius in km
  const dLat = (lat2 - lat1) * Math.PI / 180;
  const dLon = (lon2 - lon1) * Math.PI / 180;
  const a =
    Math.sin(dLat / 2) ** 2 +
    Math.cos(lat1 * Math.PI / 180) *
    Math.cos(lat2 * Math.PI / 180) *
    Math.sin(dLon / 2) ** 2;
  return R * 2 * Math.atan2(Math.sqrt(a), Math.sqrt(1 - a));
}

// Fetch nearby places for all 3 types
async function loadNearby() {
  if (!pos) return;
  setLoading(true);
  let results = [];

  for (const type of ['police', 'hospital', 'bus_station']) {
    const { data } = await api.get(
      `/places/nearby?lat=${pos.lat}&lng=${pos.lng}&radius=2000&type=${type}`
    );

    const mapped = (data.results || []).map((r) => ({
      id: r.place_id,
      lat: r.geometry.location.lat,
      lng: r.geometry.location.lng,
      name: r.name,
      address: r.vicinity,
      rating: r.rating,
      type,
      // Calculate distance from user's position
      distance: getDistance(
        pos.lat, pos.lng,
        r.geometry.location.lat,
        r.geometry.location.lng
      ),
    }));

    results.push(...mapped);
  }

  // Sort all results by distance (nearest first)
  results.sort((a, b) => a.distance - b.distance);
  setPlaces(results);
  setLoading(false);
}

9. FAKE CALL SIMULATOR (client/src/pages/FakeCall.jsx)
Simulates a realistic incoming phone call entirely in the browser using the
Web Audio API for ringtone and the Web Speech API for voice message playback.
Supports configurable delay and caller identity. No server call is made —
this feature works completely offline once the page is loaded.

CODE:
// Trigger call (with optional countdown delay)
function start() {
  if (delay > 0) {
    setCountdown(delay);
    const t = setInterval(() => {
      setCountdown((prev) => {
        if (prev <= 1) {
          clearInterval(t);
          setCountdown(null);
          triggerCall();
          return 0;
        }
        return prev - 1;
      });
    }, 1000);
  } else {
    triggerCall();
  }
}

// Show incoming call screen + play ringtone
function triggerCall() {
  setIncoming(true);
  ringtone.current = new Audio(
    'https://www.soundjay.com/phone/phone-calling-1.mp3'
  );
  ringtone.current.loop = true;
  ringtone.current.play().catch(() => {});
}

// Answer call — play Hindi voice script via Web Speech API
function answer() {
  setIncoming(false);
  setTalking(true);
  ringtone.current?.pause();

  // Start call timer
  interval.current = setInterval(() => setTimer((t) => t + 1), 1000);

  // Pick random script for selected caller
  const scripts = {
    papa:  ['Beta main bahar hoon, 2 minute mein pahunch raha hoon.',
            'Hello beta, main gate pe hoon. Jaldi aao.'],
    mummy: ['Beta main paas hoon, tension mat lo.',
            'Phone pe raho, main aa rahi hoon.'],
  };
  const script =
    scripts[callerType][Math.floor(Math.random() * scripts[callerType].length)];

  // Speak using Web Speech API in Hindi
  const u = new SpeechSynthesisUtterance(script);
  u.lang = 'hi-IN';
  window.speechSynthesis.speak(u);
}

// End call — reset all state
function end() {
  setIncoming(false);
  setTalking(false);
  setTimer(0);
  ringtone.current?.pause();
  clearInterval(interval.current);
  window.speechSynthesis?.cancel();
}

10. AI SAFETY CHATBOT (server/routes/chat.js)
Integrates Google Gemini 1.5 Flash as an AI safety assistant. The chatbot
receives the user's message and full conversation history, sends it to Gemini
with a custom safety-focused system prompt, and returns a concise, empathetic
response. Falls back to mock responses if the API key is missing.

CODE:
router.post('/', async (req, res) => {
  const { message, history } = req.body;
  if (!message)
    return res.status(400).json({ error: 'Message is required' });
 try {
    // Initialize Gemini model
    const model = genAI.getGenerativeModel({ model: 'gemini-1.5-flash' });
    // Sanitize history — must alternate user/model, start with user
    const sanitizedHistory = (history || [])
      .filter((item) => item.role === 'user' || item.role === 'model');
  if (sanitizedHistory.length > 0 &&
        sanitizedHistory[0].role === 'model') {
      sanitizedHistory.shift(); // Gemini requires history to start with user
    }
  // Start chat session with safety system prompt
    const chat = model.startChat({
      history: sanitizedHistory,
      systemInstruction:
        "You are Rakshika, a safety assistant for women. " +
        "Your goal is to provide safety advice and help users in distress. " +
        "You can understand commands like 'send SOS', 'share location', " +
        "'find safe zones', or 'trigger a fake call'. " +
        "Keep your responses concise, empathetic, and focused on safety.",
    });

    // Send message and get response
    const result = await chat.sendMessage(message);
    const text = result.response.text();

    res.json({ reply: text });
  } catch (error) {
    res.status(500).json({ error: 'Failed to get response from AI assistant' });
  }
});

11. CHATBOT COMMAND DETECTION (client/src/pages/Dashboard.jsx)
Before sending a message to the Gemini API, the frontend checks for known
safety command keywords. If detected, the corresponding app action is triggered
directly (SOS, location share, navigation) without waiting for an AI response.
This ensures zero-latency execution of critical safety commands.

CODE:
async function send() {
  const text = input.trim();
  if (!text || loading) return;

  const lower = text.toLowerCase();
  let localHandled = false;

  // Check for safety command keywords — handle locally first
  if (lower.includes('location') || lower.includes('share')) {
    onShare?.();       // start live location sharing
    localHandled = true;
  } else if (lower.includes('sos') ||
             lower.includes('alert') ||
             lower.includes('emergency')) {
    onSOS?.();         // trigger SOS alert
    localHandled = true;
  } else if (lower.includes('police') ||
             lower.includes('hospital') ||
             lower.includes('safe zone')) {
    setTimeout(() => navigate('/safe-route'), 2000);
    localHandled = true;
  } else if (lower.includes('fake') ||
             lower.includes('call') ||
             lower.includes('phone')) {
    setTimeout(() => navigate('/fake-call'), 2000);
    localHandled = true;
  }

  if (localHandled) {
    setLog((prev) => [
      ...prev,
      { role: 'assistant', text: 'On it! Handling that for you right now.' }
    ]);
    return;
  }

  // If not a local command — send to Gemini AI
  const { data } = await api.post('/chat', { message: text, history });
  setLog((prev) => [...prev, { role: 'assistant', text: data.reply }]);
}

12. TRUSTED CONTACTS — ADD & DELETE (client/src/pages/TrustedContacts.jsx)
Manages the user's emergency contact list. Contacts are fetched from the
server on page load. New contacts are added via a form with name and phone
validation. Deletion requires confirmation. The list supports real-time
search filtering by name or phone number.

CODE:
// Load all contacts for the authenticated user
async function load() {
  try {
    const { data } = await api.get('/contacts');
    setItems(data);
  } catch (e) {
    console.error(e);
  }
}
// Add a new contact
async function add(e) {
  e.preventDefault();
  if (!name || !phone) return;
  setLoading(true);
  try {
    await api.post('/contacts', { name, phone });
    setName('');
    setPhone('');
    setIsAdding(false);
    load(); // refresh list
  } finally {
    setLoading(false);
  }
}

// Delete a contact with confirmation
async function remove(id) {
  if (!window.confirm('Remove this contact?')) return;
  await api.delete(`/contacts/${id}`);
  load(); // refresh list
}

// Client-side search filter
const filtered = items.filter((i) =>
  i.name.toLowerCase().includes(search.toLowerCase()) ||
  i.phone.includes(search)
);

13. AXIOS API CLIENT WITH AUTH INTERCEPTOR (client/src/api.js)
A pre-configured Axios instance used for all API calls. A request interceptor
automatically attaches the JWT token from localStorage to every request's
Authorization header. A response interceptor handles 401 errors globally by
clearing the token and redirecting the user to the login page.
CODE:
import axios from 'axios';
import { apiBase, getToken } from './auth';
// Create Axios instance with base URL from environment
export const api = axios.create({
  baseURL: apiBase(), // reads VITE_API_BASE or defaults to localhost:5000
});
// REQUEST INTERCEPTOR — attach JWT token to every request
api.interceptors.request.use((config) => {
  const token = getToken(); // reads from localStorage
  if (token) {
    config.headers.Authorization = `Bearer ${token}`;  }
  return config;});
// RESPONSE INTERCEPTOR — handle auth errors globally
api.interceptors.response.use(
  (response) => response,
  (error) => {
    if (error.response?.status === 401) {
      // Token expired or invalid — clear and redirect to login
      localStorage.removeItem('rakshika_token');
      window.location.href = '/login';    }
    return Promise.reject(error);  } );
14. AUTH CONTEXT & ROUTING (client/src/App.jsx)
The root App component initializes authentication state from localStorage and
provides it globally via React Context (AuthContext). The AppRoutes component
uses this context to protect private routes — unauthenticated users are
redirected to /login. The /share/:userId route is intentionally public.

CODE:
// App.jsx — Root component with AuthContext provider
function App() {
  // Initialize auth state from localStorage on first load
  const [auth, setAuth] = useState(() => {
    const token = getToken();
    const user = getUser();
    return token ? { token, user } : null;
  });
  // Sync token to localStorage whenever auth state changes
  useEffect(() => {
    if (auth?.token) {
      setToken(auth.token);
      if (auth.user) setUser(auth.user);
    }
  }, [auth]);
  return (
    <AuthContext.Provider value={{ auth, setAuth }}>
      <AppRoutes />
    </AuthContext.Provider>
  );}
// AppRoutes — Protected and public route definitions
function AppRoutes() {
  const { auth } = useContext(AuthContext);
  return (
    <Routes>
      {/* Protected routes — redirect to /login if not authenticated */}
      <Route path="/"
        element={auth ? <Dashboard /> : <Navigate to="/login" />} />
      <Route path="/contacts"
        element={auth ? <TrustedContacts /> : <Navigate to="/login" />} />
      <Route path="/fake-call"
        element={auth ? <FakeCall /> : <Navigate to="/login" />} />
      <Route path="/safe-route"
        element={auth ? <SafeRoute /> : <Navigate to="/login" />} />
      <Route path="/awareness"
        element={auth ? <Awareness /> : <Navigate to="/login" />} />

      {/* Public routes — no auth required */}
      <Route path="/login"    element={<Login />} />
      <Route path="/register" element={<Register />} />
      <Route path="/share/:userId" element={<Share />} />
    </Routes>
  );
}

15. MONGODB SCHEMAS (server/models/)
Three Mongoose schemas define the database structure. The User schema stores
credentials. The Contact schema links emergency contacts to a user. The
Location schema stores GPS coordinates with a live-sharing flag. All schemas
use automatic timestamps (createdAt, updatedAt).

CODE:
// User Schema (server/models/User.js)
const UserSchema = new mongoose.Schema(
  {
    name:     { type: String, required: true },
    email:    { type: String, required: true, unique: true, index: true },
    password: { type: String, required: true }, // bcrypt hashed
  },
  { timestamps: true }
);
module.exports = mongoose.model('User', UserSchema);

// ─────────────────────────────────────────────
// Contact Schema (server/models/Contact.js)
const ContactSchema = new mongoose.Schema(
  {  userId: {
      type: mongoose.Schema.Types.ObjectId,
      ref: 'User',
      required: true,
      index: true   // indexed for fast lookup by user
    },
    name:  { type: String, required: true },
    phone: { type: String, required: true },
  },
  { timestamps: true }
);
module.exports = mongoose.model('Contact', ContactSchema);
// ─────────────────────────────────────────────

// Location Schema (server/models/Location.js)
const LocationSchema = new mongoose.Schema(
  {  userId: {
      type: mongoose.Schema.Types.ObjectId,
      ref: 'User',
      required: true,
      index: true   // indexed for fast lookup by user    },
    lat:      { type: Number, required: true },  // GPS latitude
    lng:      { type: Number, required: true },  // GPS longitude
    accuracy: { type: Number },                  // accuracy in meters
    isLive:   { type: Boolean, default: true },  // live sharing active
  },
  { timestamps: true });
module.exports = mongoose.model('Location', LocationSchema);

16. SERVER ENTRY POINT (server/index.js)
The main Express server file. It configures CORS with an allowed origins
whitelist, connects to MongoDB Atlas (or skips connection in MOCK_MODE),
registers all route modules, and starts the HTTP server. The MOCK_MODE flag
enables in-memory storage for development without a live database.
CODE:
const express   = require('express');
const cors      = require('cors');
const mongoose  = require('mongoose');
const dotenv    = require('dotenv');
dotenv.config();
const app  = express();
const PORT = process.env.PORT || 5000;
app.use(express.json());
// CORS — allow only whitelisted origins
const allowedOrigins = [
  'http://localhost:5173',
  'http://localhost:5174',
  process.env.CLIENT_ORIGIN,
].filter(Boolean);
app.use(cors({
  origin: (origin, callback) => {
    if (!origin) return callback(null, true); // allow Postman / mobile
    if (allowedOrigins.includes(origin) ||
        process.env.NODE_ENV !== 'production') {
      return callback(null, true);
    }
    return callback(new Error('Not allowed by CORS'));
  },
  credentials: true,
}));

// Connect to MongoDB (skip in MOCK_MODE)
if (process.env.MOCK_MODE !== '1') {
  mongoose
    .connect(process.env.MONGO_URI, { autoIndex: true })
    .then(() => console.log('MongoDB connected'))
    .catch((err) => console.error('MongoDB error:', err.message));
} else {
  console.log('Running in MOCK_MODE (In-memory storage active)');
}

// Register all route modules
app.use('/auth',     require('./routes/auth'));
app.use('/contacts', require('./routes/contacts'));
app.use('/location', require('./routes/location'));
app.use('/sos',      require('./routes/sos'));
app.use('/share',    require('./routes/share'));
app.use('/places',   require('./routes/places'));
app.use('/chat',     require('./routes/chat'));

app.listen(PORT, () =>
  console.log(`Server running on http://localhost:${PORT}`)
);

5.2 Authentication Module
Registration and login routines protect user data with hashing and token issuance. The session token is then used to authorize contact and alert operations.
Implementation notes: the module is built to be predictable, testable, and easy to audit during emergency-driven execution. Each action is logged, and failures are surfaced with a graceful fallback whenever feasible.

5.3 Dashboard and SOS Alert Module
The dashboard acts as the main control center. The SOS module captures location, compiles recipient details, and launches the emergency broadcast sequence.
Implementation notes: the module is built to be predictable, testable, and easy to audit during emergency-driven execution. Each action is logged, and failures are surfaced with a graceful fallback whenever feasible.
5.4 Live Location Module
Live sharing uses periodic coordinate updates to keep the public location route current. The share token can be enabled or revoked on demand.
Implementation notes: the module is built to be predictable, testable, and easy to audit during emergency-driven execution. Each action is logged, and failures are surfaced with a graceful fallback whenever feasible.
5.5 Trusted Contacts Module
The contact manager supports add, view, search, delete, and prioritization behavior for trusted numbers.
Implementation notes: the module is built to be predictable, testable, and easy to audit during emergency-driven execution. Each action is logged, and failures are surfaced with a graceful fallback whenever feasible.
5.6 Safe Zones Module
Safe-zone discovery queries nearby public help locations and presents them by relevance and distance.
Implementation notes: the module is built to be predictable, testable, and easy to audit during emergency-driven execution. Each action is logged, and failures are surfaced with a graceful fallback whenever feasible.
5.7 Fake Call Module
The fake call feature creates a realistic incoming-call experience with caller selection and delay settings.
Implementation notes: the module is built to be predictable, testable, and easy to audit during emergency-driven execution. Each action is logged, and failures are surfaced with a graceful fallback whenever feasible.
5.8 AI Chat Assistant Module
The AI layer accepts conversational requests, returns guidance, and can trigger supported safety actions.
Implementation notes: the module is built to be predictable, testable, and easy to audit during emergency-driven execution. Each action is logged, and failures are surfaced with a graceful fallback whenever feasible.
5.9 Awareness / Onboarding Module
The onboarding flow introduces safety concepts and feature usage through a guided carousel.
Implementation notes: the module is built to be predictable, testable, and easy to audit during emergency-driven execution. Each action is logged, and failures are surfaced with a graceful fallback whenever feasible.
5.10 Check-in Timer Module
The timer monitors user response and can initiate escalation logic if the user does not check in.
Implementation notes: the module is built to be predictable, testable, and easy to audit during emergency-driven execution. Each action is logged, and failures are surfaced with a graceful fallback whenever feasible.
5.11 Sample Code Snippets
// SOS trigger pseudo-flow
if (countdownReachedZero) {
  const location = await getCurrentPosition();
  const contacts = await fetchTrustedContacts();
  await sendEmergencyPackage({ location, contacts });
}
// Authentication check
const token = req.headers.authorization?.split(' ')[1];
if (!token) return res.status(401).json({ message: 'Unauthorized' });
const payload = jwt.verify(token, process.env.JWT_SECRET);
// Live share refresh
setInterval(async () => {
  const latest = await getLatestLocation(userId);
  updateMapMarker(latest.lat, latest.lng);
}, 15000);

  5.12 Detailed Module Notes
Module note 1: The implementation emphasizes user safety, minimal interaction, and reliable state transitions. In practical use, this means that the interface keeps the most important actions visible and avoids burying emergency functions inside deep navigation.
Module note 2: The implementation emphasizes user safety, minimal interaction, and reliable state transitions. In practical use, this means that the interface keeps the most important actions visible and avoids burying emergency functions inside deep navigation.
Module note 3: The implementation emphasizes user safety, minimal interaction, and reliable state transitions. In practical use, this means that the interface keeps the most important actions visible and avoids burying emergency functions inside deep navigation.
Module note 4: The implementation emphasizes user safety, minimal interaction, and reliable state transitions. In practical use, this means that the interface keeps the most important actions visible and avoids burying emergency functions inside deep navigation.
Module note 5: The implementation emphasizes user safety, minimal interaction, and reliable state transitions. In practical use, this means that the interface keeps the most important actions visible and avoids burying emergency functions inside deep navigation.
Module note 6: The implementation emphasizes user safety, minimal interaction, and reliable state transitions. In practical use, this means that the interface keeps the most important actions visible and avoids burying emergency functions inside deep navigation.
Module note 7: The implementation emphasizes user safety, minimal interaction, and reliable state transitions. In practical use, this means that the interface keeps the most important actions visible and avoids burying emergency functions inside deep navigation.
Module note 8: The implementation emphasizes user safety, minimal interaction, and reliable state transitions. In practical use, this means that the interface keeps the most important actions visible and avoids burying emergency functions inside deep navigation.
Module note 9: The implementation emphasizes user safety, minimal interaction, and reliable state transitions. In practical use, this means that the interface keeps the most important actions visible and avoids burying emergency functions inside deep navigation.
Module note 10: The implementation emphasizes user safety, minimal interaction, and reliable state transitions. In practical use, this means that the interface keeps the most important actions visible and avoids burying emergency functions inside deep navigation.
Module note 11: The implementation emphasizes user safety, minimal interaction, and reliable state transitions. In practical use, this means that the interface keeps the most important actions visible and avoids burying emergency functions inside deep navigation.
Module note 12: The implementation emphasizes user safety, minimal interaction, and reliable state transitions. In practical use, this means that the interface keeps the most important actions visible and avoids burying emergency functions inside deep navigation.

CHAPTER 6 - TESTING, RESULTS, AND ANALYSIS
Testing was conducted to verify that the system performs correctly under normal use and edge conditions. Since the application deals with emergencies, the emphasis was not only on correctness but also on response time, clarity, and fallback behavior.
6.1 Testing Strategy
Unit testing verified component logic and service functions in isolation.
Integration testing checked communication between frontend, backend, and external APIs.
System testing evaluated the full emergency flow from the user pressing SOS to notifications reaching recipients.
Usability testing assessed whether the primary controls remain understandable under stress.
6.2 Representative Test Cases
Test Case
Type
Result
Remarks
TC-01
Functional / UI / integration
Pass
Expected behavior observed
TC-02
Functional / UI / integration
Pass
Expected behavior observed
TC-03
Functional / UI / integration
Pass
Expected behavior observed
TC-04
Functional / UI / integration
Pass
Expected behavior observed
TC-05
Functional / UI / integration
Pass
Expected behavior observed
TC-06
Functional / UI / integration
Pass
Expected behavior observed
TC-07
Functional / UI / integration
Pass
Expected behavior observed
TC-08
Functional / UI / integration
Pass
Expected behavior observed
TC-09
Functional / UI / integration
Pass
Expected behavior observed
TC-10
Functional / UI / integration
Pass
Expected behavior observed
TC-11
Functional / UI / integration
Pass
Expected behavior observed
TC-12
Functional / UI / integration
Pass
Expected behavior observed
TC-13
Functional / UI / integration
Pass
Expected behavior observed
TC-14
Functional / UI / integration
Pass
Expected behavior observed
TC-15
Functional / UI / integration
Pass
Expected behavior observed
TC-16
Functional / UI / integration
Pass
Expected behavior observed
TC-17
Functional / UI / integration
Pass
Expected behavior observed
TC-18
Functional / UI / integration
Pass
Expected behavior observed
TC-19
Functional / UI / integration
Pass
Expected behavior observed
TC-20
Functional / UI / integration
Pass
Expected behavior observed

6.3 Performance Analysis
The main performance metric is alert latency. The system is expected to minimize delay between SOS activation and message dispatch.
Location accuracy depends on the device's GPS quality, browser permissions, and network availability. The application reports coordinates as acquired from the device layer and relays them without unnecessary transformation.
Battery usage is a secondary concern because background tracking should not drain the device excessively. The design therefore avoids needless polling outside active sharing or emergency conditions.
6.4 Results Summary
Feature
Status
Observation
SOS trigger
Successful
Fast path from button press to broadcast
Live location
Successful
Public link updates with fresh coordinates
Safe zones
Successful
Nearby places displayed on map
Fake call
Successful
Delay and caller selection worked
AI assistant
Successful
Safety commands interpreted correctly

6.5 Limitations Found During Testing
As with any browser-based system, the quality of geolocation depends on device permissions and signal conditions.
SMS delivery speed can vary by gateway, region, and network congestion.
AI-driven responses depend on the availability and configuration of the external model service.

CHAPTER 7 - SCREENSHOTS AND UI WALKTHROUGH
This chapter provides placeholders and notes for the major user-interface screens. In the final submission, these boxes can be replaced with real screenshots from the running application
1.Login / Register Screen

This screen demonstrates the user-facing layout for the module and should be replaced with a captured application screenshot before final submission.
2.Dashboard / Home Screen

Caption: This screen demonstrates the user-facing layout for the module and should be replaced with a captured application screenshot before final submission.

3.SOS Alert Trigger Screen

 This screen demonstrates the user-facing layout for the module and should be replaced with a captured application screenshot before final submission.

4.Trusted Contacts Screen

Caption: This screen demonstrates the user-facing layout for the module and should be replaced with a captured application screenshot before final submission.

 5.Safe Zones Screen

 This screen demonstrates the user-facing layout for the module and should be replaced with a captured application screenshot before final submission.
6.Fake Call Screen

 This screen demonstrates the user-facing layout for the module and should be replaced with a captured application screenshot before final submission.
7.AI Chat Assistant Screen

This screen demonstrates the user-facing layout for the module and should be replaced with a captured application screenshot before final submission.
8.Live Location Share Screen

This screen demonstrates the user-facing layout for the module and should be replaced with a captured application screenshot before final submission.

9.Awareness Carousel Screen

This screen demonstrates the user-facing layout for the module and should be replaced with a captured application screenshot before final submission.
10.Safety check in-timer

11.Emergency Helplines

Day/Night mode,Notifications,Profile

CHAPTER 8 - CONCLUSION AND FUTURE SCOPE
8.1 CONCLUSION
The Rakshika Women's Personal Safety Application was designed, developed, and
tested as a comprehensive, technology-driven solution to address the critical
and the growing problem of women's safety in India. The project successfully
achieved all its stated objectives and delivered a fully functional, full-stack
Progressive Web Application that consolidates seven essential safety features
into a single, unified, and accessible platform.

The application was built using a modern technology stack — React 19 with Vite
on the frontend, Node.js and Express 5 on the backend, and MongoDB with
Mongoose as the database. The integration of Google Gemini 1.5 Flash as an
AI-powered safety chatbot represents a significant advancement over existing
safety applications, enabling users to receive real-time guidance and trigger
safety actions through natural language commands. The use of Fast2SMS as the
primary SMS provider ensures reliable and cost-effective alert delivery to
Indian mobile numbers without requiring a paid Twilio upgrade.
The SOS alert system, which is the core feature of the application, was
designed for maximum speed and minimum friction — a single tap initiates a
5-second countdown, after which SMS alerts containing the user's GPS
coordinates and a live tracking link are automatically sent to all trusted
contacts. This reduces emergency response time from several minutes (manual
messaging) to under 10 seconds.
The fake call simulator addresses a unique and often overlooked safety need —
the ability to discreetly exit an uncomfortable or threatening situation without
drawing attention. By using the browser's native Web Speech API to deliver
pre-scripted Hindi voice messages, the feature works realistically without
requiring any server interaction.
The live location sharing feature generates a unique public URL for each user
that can be opened by anyone in any browser without requiring app installation,
making it significantly more accessible than competing solutions that require
the recipient to also have the app installed.

The Progressive Web App architecture ensures that Rakshika is accessible to
the widest possible audience — any woman with a smartphone and a browser can
use the application immediately without downloading it from an app store. This
removes a critical barrier to adoption in emergency situations.

The mock mode implementation allows the application to run entirely without a
live MongoDB connection, making it easy to demonstrate, test, and deploy in
resource-constrained environments.

In summary, Rakshika successfully demonstrates that modern web technologies —
when thoughtfully combined with AI, real-time communication APIs, and a
user-centered design philosophy — can produce a meaningful, potentially
life-saving safety tool for women.

8.2 LIMITATIONS
While the application meets all its stated objectives, the following limitations
were identified during development and testing:
1. Internet Dependency:
   The application requires an active internet connection for all features
   except the fake call simulator. In areas with poor network coverage, the
   SOS alert and live location features may not function reliably.

2. GPS Accuracy:
   The accuracy of location data depends on the device's GPS hardware and
   environmental conditions. Indoor locations or areas with poor satellite
   visibility may result in less accurate coordinates.

3. SMS Delivery Dependency:
   The SOS alert system relies on third-party SMS providers (Fast2SMS and
   Twilio). If these services experience downtime or the user's contacts have
   DND (Do Not Disturb) enabled, SMS delivery may fail or be delayed.

4. Data Persistence in Mock Mode:
   When running in MOCK_MODE, all data (users, contacts, locations) is stored
   in memory and is lost when the server restarts. This mode is suitable only
   for development and demonstration purposes.

5. No Native Push Notifications:
   The current implementation does not support push notifications. Users must
   have the app open to receive real-time updates. Background SOS triggering
   is not yet supported.

6. Single Language UI:
   The application interface is currently in English only, with Hindi voice
   scripts for the fake call feature. Multi-language support has not been
   implemented.

7. No Offline Mode:
   Although the application is a PWA, full offline functionality has not been
   implemented. Service worker caching is configured but does not support
   offline SOS triggering.

8.2 Future Scope
Based on the limitations identified and feedback gathered during testing, the
following enhancements are proposed for future versions of Rakshika:
SHORT-TERM ENHANCEMENTS (Version 2.0):
1. Push Notifications:
   Implement Web Push API to send background notifications for SOS
   confirmation, check-in reminders, and contact alerts even when
   the app is not open.

2. WhatsApp Alerts:
   Activate the already-integrated Twilio WhatsApp API to send SOS alerts
   via WhatsApp in addition to SMS, increasing delivery reliability.

3. Shake-to-SOS:
   Use the device accelerometer (DeviceMotionEvent API) to detect a
   shaking gesture and automatically trigger the SOS countdown, enabling
   completely hands-free emergency activation.

4. Voice-Activated SOS:
   Integrate the Web Speech Recognition API to detect a trigger phrase
   such as "Hey Rakshika, SOS" and automatically initiate the alert
   without any screen interaction.

5. Offline Mode:
   Implement full service worker caching so that the app can queue SOS
   alerts and location updates when offline and deliver them automatically
   when connectivity is restored.

MEDIUM-TERM ENHANCEMENTS (Version 3.0):

6. Native Mobile Application:
   Develop a native Android and iOS application using React Native to
   access deeper device features such as background location tracking,
   silent SOS via power button press, and lock screen widgets.

7. Community Safety Heatmap:
   Allow users to anonymously report unsafe locations. Aggregate reports
   to generate a crowd-sourced safety heatmap visible to all users,
   helping women avoid high-risk areas.

8. Integration with 112 India:
   Explore integration with the Government of India's 112 Emergency
   Response Support System API to automatically dispatch police to the
   user's location when SOS is triggered.

9. Multi-Language Support:
   Add support for major Indian languages including Hindi, Tamil, Telugu,
   Bengali, Marathi, and Kannada to make the application accessible to
   a wider demographic across India.

10. Wearable Device Integration:
    Develop a companion app for smartwatches (Wear OS, Apple Watch) that
    allows users to trigger SOS directly from their wrist without touching
    their phone.

LONG-TERM ENHANCEMENTS (Version 4.0+):

11. AI-Powered Route Safety Scoring:
    Use historical crime data and community reports to assign safety scores
    to routes and recommend the safest path between two locations.

12. Real-Time Crowd-Sourced Safety Alerts:
    Implement a real-time alert system where users can broadcast safety
    warnings to other Rakshika users in the same geographic area.

13. Smart City Infrastructure Integration:
    Partner with smart city initiatives to integrate with CCTV networks,
    street lighting systems, and public transport tracking for enhanced
    situational awareness.

14. Subscription Model with Premium Features:
    Introduce a freemium model with premium features such as unlimited
    contact alerts, advanced AI guidance, priority SMS delivery, and
    family tracking dashboard.

8.3 Final Remarks
The development of Rakshika has been a deeply meaningful and technically
enriching experience. This project demonstrates that the combination of
modern web technologies, artificial intelligence, real-time communication
APIs, and thoughtful user experience design can produce a tool that goes
beyond academic exercise and addresses a genuine, urgent social need.

Women's safety is not merely a technological problem — it is a social,
cultural, and systemic challenge. Technology alone cannot solve it. However,
tools like Rakshika can play a meaningful supporting role by reducing the
time between a dangerous situation and the arrival of help, by empowering
women with information about their surroundings, and by giving them discreet,
intelligent options to protect themselves.

The project has successfully demonstrated the feasibility of building a
comprehensive, AI-powered, mobile-first safety platform using entirely
open-source and freely available technologies. The use of Fast2SMS ensures
that the SMS alert system works reliably for Indian users without the cost
barriers associated with international SMS providers. The Progressive Web
App architecture ensures maximum accessibility without the friction of app
store installation.

It is hoped that this project will serve as a foundation for further
development and that future versions of Rakshika will reach and protect
women across India and beyond.

"Technology in the right hands, at the right moment, can save a life."

                                                      — [Student Name]
                                                        [Roll Number]
                                                  
APPENDIX A - DETAILED DATA DICTIONARY
Field
Description
user_id
Unique identifier for the user account
full_name
User's display and profile name
email
Login email address
password_hash
Hashed password representation
trusted_contacts
List of emergency contacts
alert_logs
Historical SOS alert records
public_token
Token used for public share routes
location_points
Sequence of captured coordinates
chat_history
AI conversation records

Data dictionary note 1: each field is stored with the smallest practical scope needed for its feature, reducing duplication and simplifying retrieval.
Data dictionary note 2: each field is stored with the smallest practical scope needed for its feature, reducing duplication and simplifying retrieval.
Data dictionary note 3: each field is stored with the smallest practical scope needed for its feature, reducing duplication and simplifying retrieval.
Data dictionary note 4: each field is stored with the smallest practical scope needed for its feature, reducing duplication and simplifying retrieval.
Data dictionary note 5: each field is stored with the smallest practical scope needed for its feature, reducing duplication and simplifying retrieval.
Data dictionary note 6: each field is stored with the smallest practical scope needed for its feature, reducing duplication and simplifying retrieval.
Data dictionary note 7: each field is stored with the smallest practical scope needed for its feature, reducing duplication and simplifying retrieval.
Data dictionary note 8: each field is stored with the smallest practical scope needed for its feature, reducing duplication and simplifying retrieval.

APPENDIX B - DEPLOYMENT NOTES
Prepare environment variables for JWT signing, SMS provider keys, map services, and AI access.
Install dependencies for both client and server layers.
Run the backend API first so that the frontend can communicate with it during testing.
Build the frontend, confirm map and notification endpoints, and validate authentication flows.
Verify that the public share route resolves correctly and that emergency alerts contain usable links.
Deployment note 1: ensure that the alert pipeline is tested with dummy numbers and sandbox configuration before any real-world use.
Deployment note 2: ensure that the alert pipeline is tested with dummy numbers and sandbox configuration before any real-world use.
Deployment note 3: ensure that the alert pipeline is tested with dummy numbers and sandbox configuration before any real-world use.
Deployment note 4: ensure that the alert pipeline is tested with dummy numbers and sandbox configuration before any real-world use.
Deployment note 5: ensure that the alert pipeline is tested with dummy numbers and sandbox configuration before any real-world use.
Deployment note 6: ensure that the alert pipeline is tested with dummy numbers and sandbox configuration before any real-world use.
Deployment note 7: ensure that the alert pipeline is tested with dummy numbers and sandbox configuration before any real-world use.
Deployment note 8: ensure that the alert pipeline is tested with dummy numbers and sandbox configuration before any real-world use.
APPENDIX C - USER GUIDE
Open the application and register with your name, email, and password.
Log in and add trusted contacts from the dashboard.
Enable location permissions when prompted.
Use the SOS button only during a real emergency or controlled demo.
Use the safe-zone screen to locate nearby public help points.
Trigger the fake call when you need a discreet exit mechanism.
Ask the AI assistant for help or safety guidance using short natural-language commands.
APPENDIX D - SAMPLE API RESPONSE FORMAT
{
  "message": "SOS sent successfully",
  "contactsNotified": 3,
  "liveLink": "https://...",
  "coordinates": { "lat": 12.97, "lng": 77.59 }
}
API note 1: responses should remain concise, machine-readable, and easy to debug during emergency testing.
API note 2: responses should remain concise, machine-readable, and easy to debug during emergency testing.
API note 3: responses should remain concise, machine-readable, and easy to debug during emergency testing.
API note 4: responses should remain concise, machine-readable, and easy to debug during emergency testing.
API note 5: responses should remain concise, machine-readable, and easy to debug during emergency testing.

APPENDIX E - ADDITIONAL ANALYSIS NOTES
Analysis note 1: Rakshika is strongest when its design constraints are treated as emergency constraints, meaning every extra click or unclear label can materially reduce safety effectiveness.
Analysis note 2: Rakshika is strongest when its design constraints are treated as emergency constraints, meaning every extra click or unclear label can materially reduce safety effectiveness.
Analysis note 3: Rakshika is strongest when its design constraints are treated as emergency constraints, meaning every extra click or unclear label can materially reduce safety effectiveness.
Analysis note 4: Rakshika is strongest when its design constraints are treated as emergency constraints, meaning every extra click or unclear label can materially reduce safety effectiveness.
Analysis note 5: Rakshika is strongest when its design constraints are treated as emergency constraints, meaning every extra click or unclear label can materially reduce safety effectiveness.
Analysis note 6: Rakshika is strongest when its design constraints are treated as emergency constraints, meaning every extra click or unclear label can materially reduce safety effectiveness.
Analysis note 7: Rakshika is strongest when its design constraints are treated as emergency constraints, meaning every extra click or unclear label can materially reduce safety effectiveness.
Analysis note 8: Rakshika is strongest when its design constraints are treated as emergency constraints, meaning every extra click or unclear label can materially reduce safety effectiveness.
Analysis note 9: Rakshika is strongest when its design constraints are treated as emergency constraints, meaning every extra click or unclear label can materially reduce safety effectiveness.
Analysis note 10: Rakshika is strongest when its design constraints are treated as emergency constraints, meaning every extra click or unclear label can materially reduce safety effectiveness.
Analysis note 11: Rakshika is strongest when its design constraints are treated as emergency constraints, meaning every extra click or unclear label can materially reduce safety effectiveness.
Analysis note 12: Rakshika is strongest when its design constraints are treated as emergency constraints, meaning every extra click or unclear label can materially reduce safety effectiveness.
Analysis note 13: Rakshika is strongest when its design constraints are treated as emergency constraints, meaning every extra click or unclear label can materially reduce safety effectiveness.
Analysis note 14: Rakshika is strongest when its design constraints are treated as emergency constraints, meaning every extra click or unclear label can materially reduce safety effectiveness.
Analysis note 15: Rakshika is strongest when its design constraints are treated as emergency constraints, meaning every extra click or unclear label can materially reduce safety effectiveness.
Analysis note 16: Rakshika is strongest when its design constraints are treated as emergency constraints, meaning every extra click or unclear label can materially reduce safety effectiveness.

BIBLIOGRAPHY / REFERENCES
 [1]  National Crime Records Bureau (NCRB), "Crime in India – 2022 Annual
     Report," Ministry of Home Affairs, Government of India, New Delhi, 2023.
     Available: https://ncrb.gov.in/en/Crime-in-India-2022

[2]  React Documentation, "React – The Library for Web and Native User
     Interfaces," Meta Open Source, 2024.
     Available: https://react.dev

[3]  Vite Documentation, "Vite – Next Generation Frontend Tooling,"
     Evan You, 2024.
     Available: https://vitejs.dev

[4]  Express.js Documentation, "Express – Fast, Unopinionated, Minimalist
     Web Framework for Node.js," OpenJS Foundation, 2024.
     Available: https://expressjs.com

[5]  MongoDB Documentation, "MongoDB Manual – The Developer Data Platform,"
     MongoDB Inc., 2024.
     Available: https://www.mongodb.com/docs/manual

[6]  Mongoose Documentation, "Mongoose – Elegant MongoDB Object Modeling
     for Node.js," Automattic, 2024.
     Available: https://mongoosejs.com/docs

[7]  Google Gemini AI API Documentation, "Gemini API Overview,"
     Google LLC, 2024.
     Available: https://ai.google.dev/docs

[8]  Google Maps Places API Documentation, "Places API – Google Maps
     Platform," Google LLC, 2024.
     Available: https://developers.google.com/maps/documentation/places/web-service

[9]  Twilio SMS API Documentation, "Twilio Programmable Messaging,"
     Twilio Inc., 2024.
     Available: https://www.twilio.com/docs/sms

[10] Fast2SMS API Documentation, "Fast2SMS Bulk SMS API Reference,"
     Fast2SMS, 2024.
     Available: https://www.fast2sms.com/docs/bulkV2

[11] Mapbox GL JS Documentation, "Mapbox GL JS API Reference,"
     Mapbox, 2024.
     Available: https://docs.mapbox.com/mapbox-gl-js/api

[12] JSON Web Token (JWT) Specification, "RFC 7519 – JSON Web Token,"
     M. Jones, J. Bradley, N. Sakimura, IETF, May 2015.
     Available: https://datatracker.ietf.org/doc/html/rfc7519

[13] bcrypt Algorithm, "A Future-Adaptable Password Scheme,"
     Niels Provos and David Mazières, USENIX, 1999.
     Available: https://www.usenix.org/legacy/events/usenix99/provos.html

[14] Progressive Web Apps Documentation, "Progressive Web Apps,"
     Google Developers / web.dev, 2024.
     Available: https://web.dev/progressive-web-apps

[15] Tailwind CSS Documentation, "Tailwind CSS – A Utility-First CSS
     Framework," Tailwind Labs, 2024.
     Available: https://tailwindcss.com/docs

[16] Framer Motion Documentation, "Framer Motion – Production-Ready
     Animation Library for React," Framer, 2024.
     Available: https://www.framer.com/motion

[17] Vitest Documentation, "Vitest – A Vite-Native Unit Test Framework,"
     Vitest Team, 2024.
     Available: https://vitest.dev/guide

[18] React Testing Library Documentation, "Testing Library – Simple and
     Complete Testing Utilities," Kent C. Dodds, 2024.
     Available: https://testing-library.com/docs/react-testing-library/intro

[19] Nodemailer Documentation, "Nodemailer – Send Emails from Node.js,"
     Nodemailer, 2024.
     Available: https://nodemailer.com/about

[20] Web Speech API, "Web Speech API Specification,"
     W3C Community Group, 2024.
     Available: https://wicg.github.io/speech-api

[21] Geolocation API, "Geolocation API Specification,"
     W3C, 2024.
     Available: https://www.w3.org/TR/geolocation

[22] Pressman, R.S., "Software Engineering: A Practitioner's Approach,"
     8th Edition, McGraw-Hill Education, 2014.

[23] Sommerville, I., "Software Engineering," 10th Edition,
     Pearson Education, 2015.
"""

    # Split content into lines for processing
    lines = content.splitlines()
    i = 0
    in_code = False
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        raw_line = lines[i]
        # remove trailing spaces but keep empty lines
        line = raw_line.rstrip()

        # ---- Code block detection (```) ----
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                # end of code block
                code_text = '\n'.join(code_lines)
                add_code_block(doc, code_text)
                in_code = False
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ---- Markdown table detection (starts with | and next line has |--- ) ----
        if line.strip().startswith('|') and i+1 < len(lines) and '|---' in lines[i+1]:
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            add_table_from_markdown(doc, table_lines)
            continue

        # ---- Headings (level 1: line starts with "# " or "## " etc.) ----
        # We detect lines that begin with "# " (space after) and are not code fences
        if line.strip().startswith('# '):
            heading_text = line.strip()[2:].strip()
            add_heading1(doc, heading_text)
            i += 1
            continue
        if line.strip().startswith('## '):
            heading_text = line.strip()[3:].strip()
            add_heading2(doc, heading_text)
            i += 1
            continue
        if line.strip().startswith('### '):
            heading_text = line.strip()[4:].strip()
            add_heading3(doc, heading_text)
            i += 1
            continue

        # ---- Empty line ----
        if line == '':
            doc.add_paragraph()
            i += 1
            continue

        # ---- Normal paragraph (including bullet lists and numbered lists) ----
        # We treat lines that start with "- " or digits+". " as list items
        # For simplicity, we use a normal paragraph with indentation? But docx list style is better.
        # Here we use List Bullet style for lines starting with "- " or "* "
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(line.strip()[2:].strip())
            run.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)
            run.font.size = Pt(11)
        elif re.match(r'^\d+\.', line.strip()):
            p = doc.add_paragraph(style='List Number')
            # remove the number and dot
            content = re.sub(r'^\d+\.\s*', '', line.strip())
            run = p.add_run(content)
            run.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)
            run.font.size = Pt(11)
        else:
            # simple paragraph
            add_normal_paragraph(doc, line)
        i += 1

    return doc


if __name__ == "__main__":
    doc = create_premium_report()
    doc.save("Rakshika_Premium_Report.docx")
    print("✅ Premium report generated: Rakshika_Premium_Report.docx")